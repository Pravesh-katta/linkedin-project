"""Parse resume files (DOCX, PDF, TXT) and extract skills / keywords."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
#  Common technology / skill terms to look for.
#  Kept as lower-case for case-insensitive matching.
# ---------------------------------------------------------------------------
TECH_SKILLS: set[str] = {
    # Languages
    "python", "java", "javascript", "typescript", "c", "c++", "c#",
    "go", "golang", "rust", "ruby", "php", "swift", "kotlin", "scala",
    "r", "matlab", "perl", "lua", "dart", "elixir", "haskell", "groovy",
    "objective-c", "assembly", "fortran", "cobol", "julia", "vba",
    # Web / Frontend
    "html", "css", "sass", "less", "tailwind", "bootstrap",
    "react", "reactjs", "react.js", "angular", "angularjs", "vue",
    "vuejs", "vue.js", "svelte", "next.js", "nextjs", "nuxt",
    "nuxtjs", "gatsby", "remix", "webpack", "vite", "babel",
    "jquery", "ember", "backbone",
    # Backend / Frameworks
    "node", "nodejs", "node.js", "express", "expressjs", "fastapi",
    "flask", "django", "spring", "springboot", "spring boot",
    "rails", "ruby on rails", "laravel", "asp.net", ".net",
    "dotnet", "nestjs", "koa", "hapi", "gin", "fiber", "echo",
    # Data / ML / AI
    "sql", "nosql", "mysql", "postgresql", "postgres", "mongodb",
    "redis", "elasticsearch", "cassandra", "dynamodb", "sqlite",
    "oracle", "mssql", "mariadb", "couchdb", "neo4j", "graphql",
    "pandas", "numpy", "scipy", "scikit-learn", "sklearn",
    "tensorflow", "pytorch", "keras", "opencv", "nltk", "spacy",
    "huggingface", "transformers", "llm", "gpt", "bert",
    "machine learning", "deep learning", "nlp",
    "natural language processing", "computer vision",
    "data science", "data engineering", "data analysis",
    "data analytics", "big data", "etl", "data pipeline",
    "power bi", "tableau", "looker", "spark", "pyspark",
    "hadoop", "hive", "airflow", "kafka", "flink",
    # Cloud / DevOps
    "aws", "azure", "gcp", "google cloud", "amazon web services",
    "docker", "kubernetes", "k8s", "terraform", "ansible",
    "puppet", "chef", "jenkins", "github actions", "gitlab ci",
    "circleci", "travis ci", "ci/cd", "cicd",
    "linux", "unix", "bash", "shell", "powershell",
    "nginx", "apache", "tomcat", "iis",
    "ec2", "s3", "lambda", "ecs", "eks", "fargate", "rds",
    "cloudfront", "sqs", "sns", "step functions",
    "serverless", "microservices", "monolith",
    # Tools / Misc
    "git", "github", "gitlab", "bitbucket", "svn",
    "jira", "confluence", "trello", "slack", "notion",
    "figma", "sketch", "adobe xd",
    "rest", "restful", "api", "soap", "grpc", "websocket",
    "json", "xml", "yaml", "csv", "protobuf",
    "agile", "scrum", "kanban", "waterfall",
    "tdd", "bdd", "unit testing", "integration testing",
    "selenium", "cypress", "playwright", "jest", "pytest",
    "mocha", "junit", "testng",
    # Roles / Concepts
    "full stack", "fullstack", "full-stack",
    "frontend", "front-end", "front end",
    "backend", "back-end", "back end",
    "devops", "sre", "site reliability",
    "software engineer", "software developer",
    "web developer", "mobile developer",
    "qa", "quality assurance", "automation",
    "security", "cybersecurity", "penetration testing",
    "blockchain", "web3", "solidity", "ethereum",
    "ios", "android", "react native", "flutter",
    "embedded", "iot", "firmware",
    # Soft / Process
    "leadership", "team lead", "architect",
    "project management", "product management",
    "stakeholder", "cross-functional",
    "mentoring", "code review",
}

# Multi-word skills sorted longest-first for greedy matching
_MULTI_WORD_SKILLS = sorted(
    [s for s in TECH_SKILLS if " " in s or "." in s or "/" in s],
    key=len,
    reverse=True,
)

_SINGLE_WORD_SKILLS = {s for s in TECH_SKILLS if " " not in s and "." not in s and "/" not in s}

_WORD_RE = re.compile(r"[a-zA-Z0-9#+.\-/]+")


# ---------------------------------------------------------------------------
#  File readers
# ---------------------------------------------------------------------------

def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    from docx import Document  # type: ignore[import-untyped]

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull text from tables (resumes often use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def _read_pdf(path: Path) -> str:
    from PyPDF2 import PdfReader  # type: ignore[import-untyped]

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def read_resume_file(path: Path) -> str:
    """Read a resume file and return its plain-text content."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix in {".txt", ".text", ".md"}:
        return _read_txt(path)
    raise ValueError(f"Unsupported resume format: {suffix}. Use .docx, .pdf, or .txt")


# ---------------------------------------------------------------------------
#  Keyword extraction
# ---------------------------------------------------------------------------

def extract_keywords(text: str) -> list[str]:
    """Extract technology skills and keywords from resume text.

    Returns a de-duplicated list of matched skill terms (lower-case).
    """
    lowered = text.lower()
    found: dict[str, bool] = {}

    # 1. Look for multi-word / special-char skills first
    for skill in _MULTI_WORD_SKILLS:
        if skill in lowered:
            found[skill] = True

    # 2. Look for single-word skills via word boundaries
    words = {w.lower().strip(".-") for w in _WORD_RE.findall(lowered)}
    for skill in _SINGLE_WORD_SKILLS:
        if skill in words:
            found[skill] = True

    return sorted(found.keys())


# ---------------------------------------------------------------------------
#  Resume matching
# ---------------------------------------------------------------------------

def resume_match_score(post_text: str, resume_keywords: list[str]) -> float:
    """Return 0.0–1.0 indicating how much of the resume keywords appear in
    the post text.  A score of 0.50 means 50 % keyword coverage."""
    if not resume_keywords:
        return 0.0

    lowered = post_text.lower()
    post_words = {w.lower().strip(".-") for w in _WORD_RE.findall(lowered)}

    matched = 0
    for kw in resume_keywords:
        # Multi-word keywords: substring search
        if " " in kw or "." in kw or "/" in kw:
            if kw in lowered:
                matched += 1
        else:
            if kw in post_words:
                matched += 1

    return round(matched / len(resume_keywords), 4)


def parse_and_extract(file_path: Path) -> dict[str, Any]:
    """Convenience: read file → extract text → extract keywords.

    Returns ``{"text": ..., "keywords": [...]}``.
    """
    text = read_resume_file(file_path)
    keywords = extract_keywords(text)
    return {"text": text, "keywords": keywords}
